import requests
from bs4 import BeautifulSoup
from xml.etree import ElementTree as ET
from docx import Document

# Function to parse the sitemap and extract URLs
def extract_urls_from_sitemap(sitemap_url):
    response = requests.get(sitemap_url)
    tree = ET.fromstring(response.content)
    urls = [element[0].text for element in tree if element.tag.endswith('url')]
    return urls

# Function to extract text content from a web page, excluding <nav> and <footer>
def extract_text_from_page(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    # Remove scripts, styles, <nav>, and <footer>
    for tag in soup(['script', 'style', 'nav', 'footer']):
        tag.decompose()
    # Extract text
    text = soup.get_text(separator='\n')
    # Clean up extra whitespace
    lines = (line.strip() for line in text.splitlines())
    blocks = (phrase.strip() for line in lines for phrase in line.split("  "))
    clean_text = '\n'.join(block for block in blocks if block)
    return clean_text

# Function to create a Word document with extracted text
def create_docx_document(urls, file_name):
    document = Document()
    for url in urls:
        text = extract_text_from_page(url)
        document.add_heading(url, level=1)
        document.add_paragraph(text)
    document.save(file_name)

# Sitemap URL
sitemap_url = 'http://example.com/sitemap.xml'
# Extract URLs from the sitemap
urls = extract_urls_from_sitemap(sitemap_url)
# Create a Word document with the extracted content
create_docx_document(urls, 'site_content.docx')
